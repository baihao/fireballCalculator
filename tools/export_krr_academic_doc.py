#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将 KRR_FIREBALL_PREDICTION_ACADEMIC.md 转为与 reports/总结报告.doc 风格一致的 Word 文档。"""

from __future__ import annotations

import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "source/kernel_regression/KRR_FIREBALL_PREDICTION_ACADEMIC.md"
REF_DOC = ROOT / "reports/总结报告.doc"
OUT_DOCX = ROOT / "source/kernel_regression/KRR_FIREBALL_PREDICTION_ACADEMIC.docx"
OUT_DOC = ROOT / "source/kernel_regression/KRR_FIREBALL_PREDICTION_ACADEMIC.doc"

CN = "一二三四五六七八九十"


def cn_digit(n: int) -> str:
    if n <= 10:
        return CN[n - 1]
    if n < 20:
        return "十" + CN[n - 11]
    tens, ones = divmod(n, 10)
    s = CN[tens - 1] + "十"
    if ones:
        s += CN[ones - 1]
    return s


def cn_paren(n: int) -> str:
    return f"（{cn_digit(n)}）"


def preprocess_markdown(text: str) -> str:
    lines = text.splitlines()
    out: list[str] = []
    chapter = 0
    section = 0

    for line in lines:
        if line.startswith("# ") and not line.startswith("##"):
            title = line[2:].strip()
            out.append(f"# {title}")
            chapter = 0
            section = 0
            continue
        m2 = re.match(r"^##\s+摘要\s*$", line)
        if m2:
            out.append("## 摘  要")
            chapter = 0
            section = 0
            continue
        m2 = re.match(r"^##\s+(\d+)\s+(.+)$", line)
        if m2:
            chapter = int(m2.group(1))
            section = 0
            out.append(f"## {cn_digit(chapter)}、{m2.group(2).strip()}")
            continue
        m2 = re.match(r"^##\s+(.+)$", line)
        if m2 and not line.startswith("###"):
            # 参考文献等无数字章节
            chapter += 1
            section = 0
            out.append(f"## {cn_digit(chapter)}、{m2.group(1).strip()}")
            continue
        m3 = re.match(r"^###\s+\d+\.\d+\s+(.+)$", line)
        if m3:
            section += 1
            out.append(f"### {cn_paren(section)}{m3.group(1).strip()}")
            continue
        out.append(line)
    return "\n".join(out)


def set_run_font(run, size_pt: float, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def apply_page_setup(doc: Document) -> None:
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.8)
        sec.bottom_margin = Cm(2.8)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)


def apply_paragraph_style(doc: Document) -> None:
    body_styles = {
        "Normal",
        "Body Text",
        "First Paragraph",
        "Compact",
        "List Paragraph",
    }
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        text = p.text.strip()
        if not text:
            continue
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(28)
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        if name == "Heading 1":
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            pf.left_indent = None
            for r in p.runs:
                set_run_font(r, 22, bold=True)
            continue
        if name == "Heading 2":
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = None
            for r in p.runs:
                set_run_font(r, 16, bold=True)
            continue
        if name == "Heading 3":
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = None
            for r in p.runs:
                set_run_font(r, 16, bold=False)
            continue
        if name in body_styles or name == "Normal":
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if name in {"First Paragraph", "Body Text", "Normal", "Compact"}:
                pf.first_line_indent = Pt(32)  # 约 2 字符（16pt 字号）
            for r in p.runs:
                set_run_font(r, 16, bold=False)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pf.first_line_indent = None
                    for r in p.runs:
                        set_run_font(r, 14, bold=False)


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    md_text = md_text.replace("\\Bigl[", "\\left[").replace("\\Bigr]", "\\right]")
    processed = preprocess_markdown(md_text)

    ref_docx = Path(tempfile.gettempdir()) / "总结报告_ref.docx"
    if not ref_docx.exists():
        subprocess.run(
            ["textutil", "-convert", "docx", "-output", str(ref_docx), str(REF_DOC)],
            check=True,
        )

    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", encoding="utf-8", delete=False) as f:
        f.write(processed)
        tmp_md = Path(f.name)

    subprocess.run(
        [
            "pandoc",
            str(tmp_md),
            "-o",
            str(OUT_DOCX),
            "--from",
            "markdown+tex_math_dollars+tex_math_single_backslash",
            f"--reference-doc={ref_docx}",
        ],
        check=True,
    )
    tmp_md.unlink(missing_ok=True)

    doc = Document(str(OUT_DOCX))
    apply_page_setup(doc)
    apply_paragraph_style(doc)
    doc.save(str(OUT_DOCX))

    subprocess.run(
        ["textutil", "-convert", "doc", "-output", str(OUT_DOC), str(OUT_DOCX)],
        check=True,
    )
    print(f"written {OUT_DOCX}")
    print(f"written {OUT_DOC}")


if __name__ == "__main__":
    main()
